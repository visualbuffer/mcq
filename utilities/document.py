from docx import Document
import re
from html2text import HTML2Text
from io import  BytesIO
from subprocess import check_output
from pdfminer.converter import  HTMLConverter
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfpage import PDFPage
from urllib3 import PoolManager
from pdfminer.layout import LAParams
import yaml
from shutil import copyfileobj
from datetime import datetime
import os

SETTINGS = yaml.load(open('D:\ki-ai\settings.yaml'))


class Reader(object):

    path : str = 'https://sweets.construction.com/swts_content_files_nas/153304/933113.pdf'
    text = ''
    html =''
    pages : [str] = []
    paragraphs : [str]  = []

    def __init__(self,*args, **kwagrs):
        self.text = ''
        if 'path' in kwagrs :
            self.path =  kwagrs['path']
            _file =  self.path.split('.')
            if len(_file)>1:
                http = PoolManager()
                self.response = http.request('GET', self.path, preload_content=False)
                self._file_extension  = _file[-1]
                if self._file_extension == 'docx':
                    self.docx2text()
                elif self._file_extension == 'pdf' : 
                    self.pdf2text()
                elif self._file_extension == 'doc' : 
                    self.doc2text()
                elif self._file_extension == 'rtf' : 
                    self.rtf2text()

    #TODO RTF PARSER AND DOC PARSER
    #TODO DOC PARSER

    def docx2text(self):
        document = Document(BytesIO(self.response.data))
        text = ''
        for para in document.paragraphs :
            text +=  para.text + "\n\n" 
        self.text = text
        self.preProcess()
        return self.text
    
    def doc2text(self):
        filename  =  SETTINGS['DOCUMENT']['TEMP_DIR']+str(int(datetime.now().timestamp()*1000000))+'.doc'
        antiword =  SETTINGS['DOCUMENT']['ANTIWORD']
        command = f'{antiword} {filename}'
        with open(filename, 'wb') as f:
            copyfileobj(self.response, f)
        f.close()
        text = check_output(command, shell =  True).decode()
        os.remove(filename)
        text.replace("  "," ")
        paragraphs = text.split('\r\n\r\n')
        paragraphs = [el.strip() for el in paragraphs] 
        paragraphs = [el.replace('\r\n',' ') for el in paragraphs] 
        paragraphs = [re.sub(' +',' ',el) for el in paragraphs] 
        
        self.paragraphs = [el for el in paragraphs if len(el) > 5] 
        self.text = '\n\n'.join(paragraphs)
        
        pass
    
    def rtf2text(self):
        text = BytesIO(self.response.data).read().decode()
        self.text = self._striprtf(text)
        self.preProcess()
        return self.text
        pass
    
    def text2text(self):
        self.text = BytesIO(self.response.data).read().decode()
        self.preProcess()
        return self.text


    def pdf2text(self):
        #https://www.blog.pythonlibrary.org/2018/05/03/exporting-data-from-pdfs-with-python/
        resource_manager = PDFResourceManager()
        fake_file_handle = BytesIO()
        converter = HTMLConverter(resource_manager, fake_file_handle,codec='utf-8',laparams = LAParams())
        page_interpreter = PDFPageInterpreter(resource_manager, converter)
        for page in PDFPage.get_pages(BytesIO(self.response.data), caching=True,check_extractable=True):
            page_interpreter.process_page(page)
        self.html = fake_file_handle.getvalue().decode('utf-8')
        text_maker  =  HTML2Text()
        self.text = text_maker.handle(self.html)
        self.preProcess()
        return self.text
    
    def preProcess(self):
        import re
        text =  " ".join(self.text.split('\n\n  \n\n'))
        text = text.replace("  "," ")
        page_remover = r'(.*)Page [0-9]+(.*)'
        text = re.sub( page_remover, "", text, re.M| re.I)
        self.text =  text
        paragraphs = text.split('\n\n')
        paragraphs = [el.strip() for el in paragraphs] 
        paragraphs = [el.replace('\n',' ') for el in paragraphs] 
        paragraphs = [re.sub(' +',' ',el) for el in paragraphs] 
        paragraphs = [el for el in paragraphs if len(el) > 5] 
        self.paragraphs = paragraphs
        pass




    
    def _striprtf(self, text):
       pattern = re.compile(r"\\([a-z]{1,32})(-?\d{1,10})?[ ]?|\\'([0-9a-f]{2})|\\([^a-z])|([{}])|[\r\n]+|(.)", re.I)
       # control words which specify a "destionation".
       destinations = frozenset((
          'aftncn','aftnsep','aftnsepc','annotation','atnauthor','atndate','atnicn','atnid',
          'atnparent','atnref','atntime','atrfend','atrfstart','author','background',
          'bkmkend','bkmkstart','blipuid','buptim','category','colorschememapping',
          'colortbl','comment','company','creatim','datafield','datastore','defchp','defpap',
          'do','doccomm','docvar','dptxbxtext','ebcend','ebcstart','factoidname','falt',
          'fchars','ffdeftext','ffentrymcr','ffexitmcr','ffformat','ffhelptext','ffl',
          'ffname','ffstattext','field','file','filetbl','fldinst','fldrslt','fldtype',
          'fname','fontemb','fontfile','fonttbl','footer','footerf','footerl','footerr',
          'footnote','formfield','ftncn','ftnsep','ftnsepc','g','generator','gridtbl',
          'header','headerf','headerl','headerr','hl','hlfr','hlinkbase','hlloc','hlsrc',
          'hsv','htmltag','info','keycode','keywords','latentstyles','lchars','levelnumbers',
          'leveltext','lfolevel','linkval','list','listlevel','listname','listoverride',
          'listoverridetable','listpicture','liststylename','listtable','listtext',
          'lsdlockedexcept','macc','maccPr','mailmerge','maln','malnScr','manager','margPr',
          'mbar','mbarPr','mbaseJc','mbegChr','mborderBox','mborderBoxPr','mbox','mboxPr',
          'mchr','mcount','mctrlPr','md','mdeg','mdegHide','mden','mdiff','mdPr','me',
          'mendChr','meqArr','meqArrPr','mf','mfName','mfPr','mfunc','mfuncPr','mgroupChr',
          'mgroupChrPr','mgrow','mhideBot','mhideLeft','mhideRight','mhideTop','mhtmltag',
          'mlim','mlimloc','mlimlow','mlimlowPr','mlimupp','mlimuppPr','mm','mmaddfieldname',
          'mmath','mmathPict','mmathPr','mmaxdist','mmc','mmcJc','mmconnectstr',
          'mmconnectstrdata','mmcPr','mmcs','mmdatasource','mmheadersource','mmmailsubject',
          'mmodso','mmodsofilter','mmodsofldmpdata','mmodsomappedname','mmodsoname',
          'mmodsorecipdata','mmodsosort','mmodsosrc','mmodsotable','mmodsoudl',
          'mmodsoudldata','mmodsouniquetag','mmPr','mmquery','mmr','mnary','mnaryPr',
          'mnoBreak','mnum','mobjDist','moMath','moMathPara','moMathParaPr','mopEmu',
          'mphant','mphantPr','mplcHide','mpos','mr','mrad','mradPr','mrPr','msepChr',
          'mshow','mshp','msPre','msPrePr','msSub','msSubPr','msSubSup','msSubSupPr','msSup',
          'msSupPr','mstrikeBLTR','mstrikeH','mstrikeTLBR','mstrikeV','msub','msubHide',
          'msup','msupHide','mtransp','mtype','mvertJc','mvfmf','mvfml','mvtof','mvtol',
          'mzeroAsc','mzeroDesc','mzeroWid','nesttableprops','nextfile','nonesttables',
          'objalias','objclass','objdata','object','objname','objsect','objtime','oldcprops',
          'oldpprops','oldsprops','oldtprops','oleclsid','operator','panose','password',
          'passwordhash','pgp','pgptbl','picprop','pict','pn','pnseclvl','pntext','pntxta',
          'pntxtb','printim','private','propname','protend','protstart','protusertbl','pxe',
          'result','revtbl','revtim','rsidtbl','rxe','shp','shpgrp','shpinst',
          'shppict','shprslt','shptxt','sn','sp','staticval','stylesheet','subject','sv',
          'svb','tc','template','themedata','title','txe','ud','upr','userprops',
          'wgrffmtfilter','windowcaption','writereservation','writereservhash','xe','xform',
          'xmlattrname','xmlattrvalue','xmlclose','xmlname','xmlnstbl',
          'xmlopen',
       ))
       # Translation of some special characters.
       specialchars = {
          'par': '\n',
          'sect': '\n\n',
          'page': '\n\n',
          'line': '\n',
          'tab': '\t',
          'emdash': '\u2014',
          'endash': '\u2013',
          'emspace': '\u2003',
          'enspace': '\u2002',
          'qmspace': '\u2005',
          'bullet': '\u2022',
          'lquote': '\u2018',
          'rquote': '\u2019',
          'ldblquote': '\201C',
          'rdblquote': '\u201D',
       }
       stack = []
       ignorable = False       # Whether this group (and all inside it) are "ignorable".
       ucskip = 1              # Number of ASCII characters to skip after a unicode character.
       curskip = 0             # Number of ASCII characters left to skip
       out = []                # Output buffer.
       for match in pattern.finditer(text):
          word,arg,hex,char,brace,tchar = match.groups()
          if brace:
             curskip = 0
             if brace == '{':
                # Push state
                stack.append((ucskip,ignorable))
             elif brace == '}':
                # Pop state
                ucskip,ignorable = stack.pop()
          elif char: # \x (not a letter)
             curskip = 0
             if char == '~':
                if not ignorable:
                    out.append('\xA0')
             elif char in '{}\\':
                if not ignorable:
                   out.append(char)
             elif char == '*':
                ignorable = True
          elif word: # \foo
             curskip = 0
             if word in destinations:
                ignorable = True
             elif ignorable:
                pass
             elif word in specialchars:
                out.append(specialchars[word])
             elif word == 'uc':
                ucskip = int(arg)
             elif word == 'u':
                c = int(arg)
                if c < 0: c += 0x10000
                if c > 127: out.append(chr(c)) #NOQA
                else: out.append(chr(c))
                curskip = ucskip
          elif hex: # \'xx
             if curskip > 0:
                curskip -= 1
             elif not ignorable:
                c = int(hex,16)
                if c > 127: out.append(chr(c)) #NOQA
                else: out.append(chr(c))
          elif tchar:
             if curskip > 0:
                curskip -= 1
             elif not ignorable:
                out.append(tchar)
       return ''.join(out)






